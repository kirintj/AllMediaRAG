"""文档解析器单元测试

覆盖每个 Reader 的 supported_extensions、can_handle 和基本行为。
"""

import os
import sys
import importlib
import tempfile
import pytest


# ---------------------------------------------------------------------------
# 导入辅助：确保 backend 目录在 sys.path 中
# ---------------------------------------------------------------------------

_BACKEND_DIR = os.path.normpath(os.path.join(
    os.path.dirname(__file__), "..", "..", "..", "backend",
))
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

# 确保父包存在于 sys.modules（跳过有问题的 __init__），
# 使相对导入 `from ..base import FileReader` 正常工作
if "core" not in sys.modules:
    _core_spec = importlib.util.spec_from_file_location(
        "core", os.path.join(_BACKEND_DIR, "core", "__init__.py"),
        submodule_search_locations=[os.path.join(_BACKEND_DIR, "core")],
    )
    _core_mod = importlib.util.module_from_spec(_core_spec)
    sys.modules["core"] = _core_mod
    _core_spec.loader.exec_module(_core_mod)

if "core.providers" not in sys.modules:
    _providers_spec = importlib.util.spec_from_file_location(
        "core.providers",
        os.path.join(_BACKEND_DIR, "core", "providers", "__init__.py"),
        submodule_search_locations=[os.path.join(_BACKEND_DIR, "core", "providers")],
    )
    _providers_mod = importlib.util.module_from_spec(_providers_spec)
    sys.modules["core.providers"] = _providers_mod

    # 先加载 base 模块，再注册到 providers 包
    _base_path = os.path.join(_BACKEND_DIR, "core", "providers", "base.py")
    _base_spec = importlib.util.spec_from_file_location("core.providers.base", _base_path)
    _base_mod = importlib.util.module_from_spec(_base_spec)
    sys.modules["core.providers.base"] = _base_mod
    _base_spec.loader.exec_module(_base_mod)
    _providers_mod.base = _base_mod

if "core.providers.readers" not in sys.modules:
    _readers_spec = importlib.util.spec_from_file_location(
        "core.providers.readers",
        os.path.join(_BACKEND_DIR, "core", "providers", "readers", "__init__.py"),
        submodule_search_locations=[os.path.join(_BACKEND_DIR, "core", "providers", "readers")],
    )
    _readers_mod = importlib.util.module_from_spec(_readers_spec)
    sys.modules["core.providers.readers"] = _readers_mod


def _load_reader_module(module_name: str):
    """加载 readers 子模块（父包已预注册，相对导入可正常工作）"""
    full_name = f"core.providers.readers.{module_name}"
    if full_name in sys.modules:
        return sys.modules[full_name]

    module_path = os.path.join(
        _BACKEND_DIR, "core", "providers", "readers", f"{module_name}.py",
    )
    spec = importlib.util.spec_from_file_location(full_name, module_path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[full_name] = module
    spec.loader.exec_module(module)
    return module


def _get_pdf_reader_class():
    mod = _load_reader_module("pdf_reader")
    return mod.PDFReader


def _get_markdown_reader_class():
    mod = _load_reader_module("markdown_reader")
    return mod.MarkdownReader


def _get_docx_reader_class():
    mod = _load_reader_module("docx_reader")
    return mod.DocxReader


def _get_html_reader_class():
    mod = _load_reader_module("html_reader")
    return mod.HtmlReader


def _get_image_reader_class():
    mod = _load_reader_module("image_reader")
    return mod.ImageReader


# ---------------------------------------------------------------------------
# PDFReader
# ---------------------------------------------------------------------------

class TestPDFReader:
    """PDFReader 测试"""

    def test_supported_extensions(self):
        PDFReader = _get_pdf_reader_class()
        reader = PDFReader()
        assert reader.supported_extensions() == [".pdf"]

    def test_can_handle(self):
        PDFReader = _get_pdf_reader_class()
        reader = PDFReader()
        assert reader.can_handle("report.pdf") is True
        assert reader.can_handle("report.PDF") is True
        assert reader.can_handle("report.txt") is False

    def test_requires_pypdf2(self, monkeypatch):
        """缺少 PyPDF2 时应抛出 ImportError"""
        import builtins
        original_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "PyPDF2":
                raise ImportError("No module named 'PyPDF2'")
            return original_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)

        PDFReader = _get_pdf_reader_class()
        reader = PDFReader()
        with pytest.raises(ImportError, match="PyPDF2"):
            reader.read("dummy.pdf")


# ---------------------------------------------------------------------------
# MarkdownReader
# ---------------------------------------------------------------------------

class TestMarkdownReader:
    """MarkdownReader 测试"""

    def test_supported_extensions(self):
        MarkdownReader = _get_markdown_reader_class()
        reader = MarkdownReader()
        exts = reader.supported_extensions()
        assert ".md" in exts
        assert ".markdown" in exts

    def test_can_handle(self):
        MarkdownReader = _get_markdown_reader_class()
        reader = MarkdownReader()
        assert reader.can_handle("readme.md") is True
        assert reader.can_handle("notes.markdown") is True
        assert reader.can_handle("notes.txt") is False

    def test_read_md_file(self, tmp_path):
        """读取临时 Markdown 文件并检查输出包含 HTML"""
        MarkdownReader = _get_markdown_reader_class()

        md_content = "# Hello\n\nThis is **bold** text.\n"
        md_file = tmp_path / "test.md"
        md_file.write_text(md_content, encoding="utf-8")

        reader = MarkdownReader()
        result = reader.read(str(md_file))

        # markdown 库输出包含 HTML 标签
        assert "Hello" in result
        assert "bold" in result

    def test_read_fenced_code(self, tmp_path):
        """Markdown 代码块应保留"""
        MarkdownReader = _get_markdown_reader_class()

        md_content = "```python\nprint('hello')\n```\n"
        md_file = tmp_path / "code.md"
        md_file.write_text(md_content, encoding="utf-8")

        reader = MarkdownReader()
        result = reader.read(str(md_file))

        # fenced_code 扩展生成 <code> 标签，内容中包含 print
        assert "print" in result
        assert "hello" in result


# ---------------------------------------------------------------------------
# DocxReader
# ---------------------------------------------------------------------------

class TestDocxReader:
    """DocxReader 测试"""

    def test_supported_extensions(self):
        DocxReader = _get_docx_reader_class()
        reader = DocxReader()
        assert reader.supported_extensions() == [".docx"]

    def test_can_handle(self):
        DocxReader = _get_docx_reader_class()
        reader = DocxReader()
        assert reader.can_handle("doc.docx") is True
        assert reader.can_handle("doc.pdf") is False

    def test_read_docx_file(self, tmp_path):
        """读取临时 docx 文件"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装")

        DocxReader = _get_docx_reader_class()

        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        reader = DocxReader()
        result = reader.read(str(docx_path))

        assert "第一段内容" in result
        assert "第二段内容" in result


# ---------------------------------------------------------------------------
# HtmlReader
# ---------------------------------------------------------------------------

class TestHtmlReader:
    """HtmlReader 测试"""

    def test_supported_extensions(self):
        HtmlReader = _get_html_reader_class()
        reader = HtmlReader()
        exts = reader.supported_extensions()
        assert ".html" in exts
        assert ".htm" in exts

    def test_can_handle(self):
        HtmlReader = _get_html_reader_class()
        reader = HtmlReader()
        assert reader.can_handle("page.html") is True
        assert reader.can_handle("page.htm") is True
        assert reader.can_handle("page.txt") is False

    def test_strips_nav_footer_script(self, tmp_path):
        """nav、footer、script 标签应被移除"""
        HtmlReader = _get_html_reader_class()

        html = (
            "<html><body>"
            "<nav>Navigation Menu</nav>"
            "<main><p>Main content here</p></main>"
            "<footer>Footer text</footer>"
            "<script>alert('xss')</script>"
            "</body></html>"
        )
        html_file = tmp_path / "test.html"
        html_file.write_text(html, encoding="utf-8")

        reader = HtmlReader()
        result = reader.read(str(html_file))

        assert "Main content here" in result
        assert "Navigation Menu" not in result
        assert "Footer text" not in result
        assert "alert" not in result

    def test_extracts_main_content(self, tmp_path):
        """优先从 <main> 或 <article> 提取内容"""
        HtmlReader = _get_html_reader_class()

        html = (
            "<html><body>"
            "<div>Sidebar stuff</div>"
            "<article><h1>Article Title</h1><p>Article body.</p></article>"
            "</body></html>"
        )
        html_file = tmp_path / "article.html"
        html_file.write_text(html, encoding="utf-8")

        reader = HtmlReader()
        result = reader.read(str(html_file))

        assert "Article Title" in result
        assert "Article body." in result

    def test_preserves_code_blocks(self, tmp_path):
        """代码块应保留并用 ``` 标记"""
        HtmlReader = _get_html_reader_class()

        html = (
            "<html><body>"
            "<main><pre><code>def hello():\n    print('hi')</code></pre></main>"
            "</body></html>"
        )
        html_file = tmp_path / "code.html"
        html_file.write_text(html, encoding="utf-8")

        reader = HtmlReader()
        result = reader.read(str(html_file))

        assert "def hello():" in result
        assert "```" in result


# ---------------------------------------------------------------------------
# ImageReader
# ---------------------------------------------------------------------------

class TestImageReader:
    """ImageReader 测试"""

    def test_supported_extensions(self):
        ImageReader = _get_image_reader_class()
        reader = ImageReader()
        exts = reader.supported_extensions()
        assert ".png" in exts
        assert ".jpg" in exts
        assert ".jpeg" in exts
        assert ".bmp" in exts
        assert ".tiff" in exts
        assert ".tif" in exts

    def test_can_handle(self):
        ImageReader = _get_image_reader_class()
        reader = ImageReader()
        assert reader.can_handle("photo.png") is True
        assert reader.can_handle("scan.tiff") is True
        assert reader.can_handle("doc.pdf") is False

    def test_no_providers_returns_empty(self, tmp_path):
        """未配置任何 provider 时返回空字符串"""
        ImageReader = _get_image_reader_class()

        png_path = tmp_path / "test.png"
        png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        reader = ImageReader()
        result = reader.read(str(png_path))

        assert result == ""

    def test_file_not_found(self):
        """文件不存在时抛出 FileNotFoundError"""
        ImageReader = _get_image_reader_class()

        reader = ImageReader()
        with pytest.raises(FileNotFoundError):
            reader.read("/nonexistent/image.png")

    def test_with_mock_ocr(self, tmp_path):
        """配置 mock OCR provider 时提取文字"""
        ImageReader = _get_image_reader_class()

        png_path = tmp_path / "test.png"
        png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        class MockOCR:
            def extract_text(self, path):
                return "extracted text from image"

        reader = ImageReader(ocr_provider=MockOCR())
        result = reader.read(str(png_path))

        assert "extracted text from image" in result
        assert "[OCR 文字内容]" in result

    def test_with_mock_vlm(self, tmp_path):
        """配置 mock VLM provider 时生成描述"""
        ImageReader = _get_image_reader_class()

        png_path = tmp_path / "test.png"
        png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        class MockVLM:
            def describe_image(self, path):
                return "这是一张测试图片"

        reader = ImageReader(vlm_provider=MockVLM())
        result = reader.read(str(png_path))

        assert "这是一张测试图片" in result
        assert "[图片内容描述]" in result

    def test_with_both_providers(self, tmp_path):
        """同时配置 OCR 和 VLM 时两者结果均包含"""
        ImageReader = _get_image_reader_class()

        png_path = tmp_path / "test.png"
        png_path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        class MockOCR:
            def extract_text(self, path):
                return "图片中的文字"

        class MockVLM:
            def describe_image(self, path):
                return "这是一张流程图"

        reader = ImageReader(ocr_provider=MockOCR(), vlm_provider=MockVLM())
        result = reader.read(str(png_path))

        assert "[OCR 文字内容]" in result
        assert "图片中的文字" in result
        assert "[图片内容描述]" in result
        assert "这是一张流程图" in result


# ---------------------------------------------------------------------------
# 包导入测试
# ---------------------------------------------------------------------------

class TestPackageImports:
    """验证 readers 包可正常导入"""

    def test_import_all_readers(self):
        PDFReader = _get_pdf_reader_class()
        MarkdownReader = _get_markdown_reader_class()
        DocxReader = _get_docx_reader_class()
        HtmlReader = _get_html_reader_class()
        ImageReader = _get_image_reader_class()

        assert PDFReader is not None
        assert MarkdownReader is not None
        assert DocxReader is not None
        assert HtmlReader is not None
        assert ImageReader is not None

    def test_all_reader_subclass_of_file_reader(self):
        """所有 Reader 都是 FileReader 子类"""
        FileReader = sys.modules["core.providers.base"].FileReader

        PDFReader = _get_pdf_reader_class()
        MarkdownReader = _get_markdown_reader_class()
        DocxReader = _get_docx_reader_class()
        HtmlReader = _get_html_reader_class()
        ImageReader = _get_image_reader_class()

        for cls in (PDFReader, MarkdownReader, DocxReader, HtmlReader, ImageReader):
            assert issubclass(cls, FileReader), f"{cls.__name__} 不是 FileReader 子类"
